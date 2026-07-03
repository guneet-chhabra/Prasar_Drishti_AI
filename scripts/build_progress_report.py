from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_PATH = OUT_DIR / "Prasar_Drishti_AI_Report_Till_Now.docx"


def read_text(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8")


def safe_json(path: str) -> list[dict]:
    file_path = ROOT / path
    if not file_path.exists():
        return []
    return json.loads(file_path.read_text(encoding="utf-8"))


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.1)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(code.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(8)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    doc.add_paragraph()


def file_tree() -> str:
    paths = [
        "backend/run.py",
        "backend/config.py",
        "backend/app/__init__.py",
        "backend/api/auth.py",
        "backend/api/news.py",
        "backend/scraper/newsonair_scraper.py",
        "backend/scripts/run_news_monitor.py",
        "backend/ml/MODEL_PLAN.md",
        "backend/data/raw/newsonair_articles.json",
        "backend/data/raw/today_YYYY-MM-DD.json",
        "backend/data/raw/archive_articles.json",
        "frontend/index.html",
        "frontend/app/auth.js",
        "frontend/styles/auth.css",
        "README.md",
    ]
    return "\n".join(paths)


def snippet(text: str, start: str, end: str | None = None, max_lines: int = 35) -> str:
    lines = text.splitlines()
    start_index = next((i for i, line in enumerate(lines) if start in line), 0)
    if end:
        end_index = next((i for i, line in enumerate(lines[start_index:], start_index) if end in line), start_index + max_lines)
    else:
        end_index = start_index + max_lines
    return "\n".join(lines[start_index : min(end_index + 1, len(lines))])


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Prasar Drishti AI\nProject Progress Report")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(3, 20, 39)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Report generated on {date.today().isoformat()}").italic = True
    doc.add_paragraph()

    add_heading(doc, "1. Project Overview", 1)
    add_body(
        doc,
        "Prasar Drishti AI is an AI/ML-based Sports Intelligence and News Analytics System. "
        "The project collects NewsOnAir articles, stores them as structured datasets, and will use them for "
        "sentiment analysis, Indian sports forecasting, and FIFA/football prediction features.",
    )
    add_bullets(
        doc,
        [
            "Completed so far: login authentication with admin/user roles.",
            "Completed so far: NewsOnAir scraper with archive, today, and monitor workflow.",
            "Completed so far: frontend controls for login, article display, daily update, archive view, and admin actions.",
            "Planned next: real database tables, labelled datasets, model training, dashboards, and prediction APIs.",
        ],
    )

    add_heading(doc, "2. Current File Structure", 1)
    add_code(doc, file_tree())

    add_heading(doc, "3. Authentication Module", 1)
    add_body(
        doc,
        "The backend uses Flask routes for login, logout, current-user lookup, and admin-only access. "
        "For local development reliability, the backend returns a signed bearer token at login. The frontend stores "
        "that token in localStorage and sends it with API calls using the Authorization header.",
    )
    add_table(
        doc,
        ["Route", "Method", "Access", "Purpose"],
        [
            ["/api/auth/login", "POST", "Public", "Validates username/password and returns user + token."],
            ["/api/auth/me", "GET", "Logged-in user", "Returns current authenticated user."],
            ["/api/auth/logout", "POST", "Logged-in user", "Clears Flask session."],
            ["/api/auth/admin/overview", "GET", "Admin", "Checks role-based admin access."],
        ],
    )

    auth_code = read_text("backend/api/auth.py")
    add_heading(doc, "Authentication Code Snippet", 2)
    add_code(doc, snippet(auth_code, "def create_token", "def login_required", max_lines=45))
    add_code(doc, snippet(auth_code, "@auth_bp.post", "@auth_bp.get(\"/me\")", max_lines=35))

    add_heading(doc, "4. Frontend Website Work", 1)
    add_body(
        doc,
        "The frontend is currently a static HTML/CSS/JavaScript interface. It contains a dark command-center style "
        "login page inspired by the supplied UI idea, then reveals dashboard controls after successful login.",
    )
    add_table(
        doc,
        ["Control", "Visible To", "Purpose"],
        [
            ["Today", "Admin and user", "Loads today's rolling scraped dataset."],
            ["Archive", "Admin and user", "Loads long-term historical dataset."],
            ["Update Today", "Admin", "Runs a scrape for today's articles."],
            ["Scrape Last Month", "Admin", "Builds initial historical archive."],
            ["Archive Today", "Admin", "Merges today's data into archive."],
        ],
    )

    index_code = read_text("frontend/index.html")
    js_code = read_text("frontend/app/auth.js")
    add_heading(doc, "Website HTML Snippet", 2)
    add_code(doc, snippet(index_code, "<section class=\"login-panel\"", "</section>", max_lines=45))
    add_heading(doc, "Frontend API Snippet", 2)
    add_code(doc, snippet(js_code, "async function api", "loginForm.addEventListener", max_lines=35))

    add_heading(doc, "5. NewsOnAir Scraper And Data Pipeline", 1)
    add_body(
        doc,
        "The scraper scans NewsOnAir category pages, follows article links, parses article title/date/body, marks "
        "government-related content, and saves deduplicated JSON/CSV files. It now collects broad categories, not just "
        "government content, so the same archive can support sentiment analysis, Indian sports prediction, and FIFA-related analysis.",
    )
    add_table(
        doc,
        ["Dataset", "File Pattern", "Purpose"],
        [
            ["Today", "backend/data/raw/today_YYYY-MM-DD.json", "Rolling daily dataset updated every 30 minutes."],
            ["Archive", "backend/data/raw/archive_articles.json", "Long-term historical corpus for ML training and analysis."],
            ["Default/demo", "backend/data/raw/newsonair_articles.json", "Starter sample dataset used during early UI testing."],
        ],
    )

    scraper_code = read_text("backend/scraper/newsonair_scraper.py")
    news_code = read_text("backend/api/news.py")
    monitor_code = read_text("backend/scripts/run_news_monitor.py")
    add_heading(doc, "Scraper Code Snippet", 2)
    add_code(doc, snippet(scraper_code, "def scrape(", "def save(", max_lines=50))
    add_heading(doc, "News API Snippet", 2)
    add_code(doc, snippet(news_code, "@news_bp.get(\"/articles\")", "@news_bp.post(\"/scrape-last-month\")", max_lines=65))
    add_heading(doc, "Half-Hour Monitor Snippet", 2)
    add_code(doc, snippet(monitor_code, "def update_today", "def main", max_lines=45))

    add_heading(doc, "6. Sample Scraped Data", 1)
    articles = safe_json("backend/data/raw/newsonair_articles.json")[:3]
    rows = []
    for article in articles:
        rows.append(
            [
                article.get("title", "")[:80],
                article.get("category", ""),
                article.get("published_at", ""),
                str(article.get("word_count", "")),
            ]
        )
    add_table(doc, ["Title", "Category", "Published At", "Words"], rows)
    add_heading(doc, "JSON Data Snippet", 2)
    add_code(doc, json.dumps(articles[:1], indent=2, ensure_ascii=False))

    add_heading(doc, "7. Proposed Database Schema", 1)
    add_body(
        doc,
        "The project currently stores data as JSON/CSV for quick learning and debugging. The next backend step should "
        "move users and articles into a relational database. The following schema is recommended.",
    )
    add_code(
        doc,
        """
CREATE TABLE users (
    id SERIAL PRIMARY KEY,
    username VARCHAR(80) UNIQUE NOT NULL,
    password_hash TEXT NOT NULL,
    role VARCHAR(20) NOT NULL CHECK (role IN ('admin', 'user')),
    display_name VARCHAR(120),
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE articles (
    id SERIAL PRIMARY KEY,
    url TEXT UNIQUE NOT NULL,
    title TEXT NOT NULL,
    category VARCHAR(40),
    published_at VARCHAR(80),
    published_date DATE,
    summary TEXT,
    body TEXT,
    word_count INTEGER,
    is_government_related BOOLEAN DEFAULT FALSE,
    scraped_at TIMESTAMP,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE model_runs (
    id SERIAL PRIMARY KEY,
    model_name VARCHAR(120) NOT NULL,
    task_name VARCHAR(120) NOT NULL,
    accuracy NUMERIC(6, 4),
    macro_f1 NUMERIC(6, 4),
    notes TEXT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);
""",
    )

    add_heading(doc, "8. ML Model Plan", 1)
    add_body(
        doc,
        "The first ML task should be sentiment/orientation classification, followed by sports-specific prediction features. "
        "The model strategy is to start with simple baselines, then move to heavier neural models only after a clean labelled dataset exists.",
    )
    add_table(
        doc,
        ["Model", "Use Case", "Notes"],
        [
            ["TF-IDF + Logistic Regression", "Sentiment baseline", "Fast, explainable, good first benchmark."],
            ["TF-IDF + Linear SVM", "Sentiment baseline", "Often strong for text classification."],
            ["Naive Bayes", "Light baseline", "Simple and useful for comparison."],
            ["Random Forest / XGBoost", "Feature-based prediction", "Useful for structured sports indicators."],
            ["MLP", "Neural baseline", "Use ReLU hidden layers and Softmax output."],
            ["LSTM / BiLSTM", "Sequence model", "Uses Tanh/Sigmoid internally; needs more data."],
            ["DistilBERT", "Transformer classifier", "Uses GELU internally; best after enough labels."],
        ],
    )

    add_heading(doc, "9. Current Run Commands", 1)
    add_code(
        doc,
        """
cd backend
.\\venv\\Scripts\\activate
python -m pip install -r requirements-scraper.txt
python run.py

cd frontend
python -m http.server 3000

python -m scraper.newsonair_scraper --days 30 --pages 35 --limit 300 --stem archive_articles
python -m scraper.newsonair_scraper --today --pages 3 --limit 40
python scripts\\run_news_monitor.py
""",
    )

    add_heading(doc, "10. Next Steps", 1)
    add_bullets(
        doc,
        [
            "Move demo users into a real database table.",
            "Move article JSON/CSV storage into database tables while keeping CSV export.",
            "Create labelled sentiment dataset from archived articles.",
            "Train and compare baseline ML models using Accuracy and Macro F1-score.",
            "Build charts for sentiment trends, sports trends, and FIFA-related article tracking.",
        ],
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
