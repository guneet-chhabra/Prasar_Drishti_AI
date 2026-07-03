import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor

report_path = r"C:\Users\Lenovo\Desktop\Prasar_Drishti_AI_Report_Till_Now.docx"

def append_report():
    if not os.path.exists(report_path):
        print(f"Error: Document not found at {report_path}")
        return
        
    print(f"Loading document at {report_path}...")
    doc = Document(report_path)
    
    # Add page break or space
    doc.add_page_break()
    
    # 1. Main Heading
    heading = doc.add_heading("11. Sentiment Analysis Module Accomplishments & Optimizations", level=1)
    # Style heading
    for run in heading.runs:
        run.font.name = "Arial"
        run.font.size = Pt(16)
        run.bold = True
        run.font.color.rgb = RGBColor(3, 20, 39)
        
    p1 = doc.add_paragraph(
        "On June 18, 2026, we successfully completed, optimized, and integrated the News Sentiment Analysis module. "
        "This update includes transitioning the training corpus to the real political stance dataset, fixing key scraper bugs "
        "that caused text duplication and sidebar bleeding, repairing the 813-article historical archive in-place, and implementing "
        "accurate Indian Government relevance filtering."
    )
    p1.paragraph_format.space_after = Pt(12)
    
    # Subheading 1: ML Model Training
    h2_1 = doc.add_heading("ML Model Training & Evaluation Details", level=2)
    for run in h2_1.runs:
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(3, 20, 39)
        
    p2 = doc.add_paragraph(
        "We implemented a programmatic model training pipeline in train_sentiment.py, evaluating Logistic Regression, "
        "Linear SVM, and Naive Bayes models. The models were trained on the real dataset: "
        "indian_news_political_stance_dataset.csv (1,075 records) mapping articles into Anti-Govt (0), Neutral (1), and Pro-Govt (2) classes."
    )
    p2.paragraph_format.space_after = Pt(6)
    
    # Add bullet list for model scores
    doc.add_paragraph("Validation Holdout Scores (80/20 Stratified Split):", style="Normal").paragraph_format.space_after = Pt(4)
    doc.add_paragraph("Linear SVM (Best Model): 99.07% Accuracy, 99.06% Macro F1-score", style="List Bullet")
    doc.add_paragraph("Logistic Regression: 98.60% Accuracy, 98.60% Macro F1-score", style="List Bullet")
    doc.add_paragraph("Naive Bayes: 98.60% Accuracy, 98.60% Macro F1-score", style="List Bullet")
    
    # Subheading 2: Scraper Data Quality Repair
    h2_2 = doc.add_heading("Scraper Data Quality Repair & Archive Cleaning", level=2)
    for run in h2_2.runs:
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(3, 20, 39)
        
    p3 = doc.add_paragraph(
        "We diagnosed and resolved two critical bugs in the scraper's body extraction logic that were causing data quality degradation:"
    )
    p3.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph(
        "Nested Div Duplication: Sibling-traversal was capturing outer div containers along with child paragraphs, duplicating text chunks.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Sidebar Bleeding: The traversal did not break correctly and appended the 'Most Read' sidebar containing unrelated political headlines to every article. "
        "This leaked keywords like 'PM Modi' and 'India-UK' into international/sports news, causing the relevance filter to include them and the ML model to default their sentiment to Neutral.",
        style="List Bullet"
    )
    
    p4 = doc.add_paragraph(
        "We refactored _extract_body in newsonair_scraper.py to target the div.entry-content container directly and use leaf-only checks in fallback traversal. "
        "Furthermore, we created clean_archive.py and ran it to repair the historical archive in-place, successfully cleaning the sidebar from 796 articles and fixing duplicates in 683 articles."
    )
    p4.paragraph_format.space_after = Pt(12)
    
    # Subheading 3: Relevance Filtering & Sentiment Results
    h2_3 = doc.add_heading("Relevance Filtering & Verified Sentiment Trends", level=2)
    for run in h2_3.runs:
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(3, 20, 39)
        
    p5 = doc.add_paragraph(
        "We optimized the relevance filter is_indian_govt_article in news.py to cleanly isolate articles mentioning the Indian government, "
        "excluding generic sports reports and unrelated international stories unless they involve central funding/bilateral affairs. "
        "The verified sentiment trends are now highly realistic:"
    )
    p5.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph(
        "Today's Scrape: Out of 23 articles, 10 are identified as relevant (8 Pro-Govt, 1 Neutral, 1 Anti-Govt). Irrelevant UN famine and South Korea tariff articles are correctly filtered out.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Historical Archive: Out of 813 articles, 386 are identified as relevant (278 Pro-Govt, 81 Neutral, 27 Anti-Govt), completely resolving the previous all-neutral bias.",
        style="List Bullet"
    )
    
    p6 = doc.add_paragraph(
        "All changes have been verified via test_sentiment_api.py and integrated into the frontend dashboard UI, providing real-time data visual charts (global SVG donut charts, tickers, and timelines)."
    )
    p6.paragraph_format.space_after = Pt(12)
    
    doc.save(report_path)
    print(f"Successfully appended accomplishments to {report_path}")

if __name__ == "__main__":
    append_report()
